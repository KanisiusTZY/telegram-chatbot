"""
Tool definitions and implementations for the Telegram AI agent.

Each tool has:
  - A JSON schema (TOOLS list) — passed to Groq's tools= parameter
  - An implementation function — called when the model requests it
"""

import json
import logging
import base64
from datetime import datetime, timezone

from simpleeval import simple_eval, EvalWithCompoundTypes, InvalidExpression
from duckduckgo_search import DDGS

import agent_db

log = logging.getLogger(__name__)

# ─── Tool Schemas (OpenAI / Groq format) ─────────────────────────────────────

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Cari informasi real-time dan fakta di internet pakai DuckDuckGo. "
                "WAJIB digunakan saat user menanyakan pertanyaan tentang istilah, singkatan, tim esport/olahraga (misal BTR, RRQ, EVOS), "
                "tokoh, berita, fakta, harga, cuaca, atau topik umum apa pun agar jawaban selalu akurat dan terbaru."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Query pencarian dalam bahasa Indonesia atau Inggris (misal 'apa itu BTR esport', 'harga btc hari ini').",
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_reminder",
            "description": (
                "WAJIB panggil tool ini setiap kali user minta pengingat/reminder dalam durasi atau waktu apapun "
                "(misal '2 detik', '5 menit', '1 jam', 'besok jam 9'). "
                "bisa menerima string waktu relatif seperti '2 detik', '5 menit', '+10s' atau ISO format."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {
                        "type": "string",
                        "description": "Isi pesan pengingat.",
                    },
                    "remind_at": {
                        "type": "string",
                        "description": "Waktu pengingat (contoh: '2 detik', '5 menit', '1 jam', '+30s', atau ISO format).",
                    },
                },
                "required": ["message", "remind_at"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_note",
            "description": "Simpan catatan personal untuk user. Cocok buat nyimpen info penting.",
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "Isi catatan yang mau disimpan.",
                    },
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_notes",
            "description": "Ambil semua catatan personal milik user.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate",
            "description": (
                "Hitung ekspresi matematika dengan aman. "
                "Contoh: '2 ** 10', '(3.14 * 5**2)', 'sqrt(144)' (butuh 'math.' prefix: 'math.sqrt(144)'). "
                "Jangan pakai buat hal di luar matematika."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "expression": {
                        "type": "string",
                        "description": "Ekspresi matematika yang mau dihitung.",
                    },
                },
                "required": ["expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "file_convert",
            "description": (
                "Konversi file yang baru saja dikirim user ke format lain. "
                "Gunakan setelah user mengirim file dan meminta konversi (misal 'ubah ke PDF', 'jadiin docx', 'jadiin png', 'ekstrak ke txt')."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "target_format": {
                        "type": "string",
                        "enum": ["pdf", "docx", "png", "jpg", "txt"],
                        "description": "Format target konversi.",
                    },
                },
                "required": ["target_format"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "media_download",
            "description": (
                "Unduh video atau MP3 audio dari link sosial media (TikTok, Instagram Reels, YouTube Shorts, Twitter/X, dll). "
                "Gunakan saat user mengirimkan link media sosial atau meminta unduh video/audio dari link."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "URL/link media sosial yang mau diunduh.",
                    },
                    "extract_audio": {
                        "type": "boolean",
                        "description": "Set true jika user minta format MP3 / audio saja.",
                    },
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "remove_bg",
            "description": (
                "Hapus background dari foto/gambar yang diunggah user. Bisa ubah ke transparan (PNG) atau pasfoto background merah/biru/putih. "
                "Gunakan saat user minta hapus background foto, bikin PNG transparan, atau ubah warna background pasfoto."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "bg_color": {
                        "type": "string",
                        "enum": ["transparent", "merah", "biru", "putih"],
                        "description": "Warna background target: 'transparent' (default PNG), 'merah', 'biru', atau 'putih'.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "music_search",
            "description": (
                "Cari dan unduh lagu/musik MP3 berdasarkan judul lagu atau nama penyanyi. "
                "Gunakan saat user minta 'cari lagu [judul]', 'download mp3 [judul]', atau 'setel lagu [judul]'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Judul lagu atau nama penyanyi yang ingin dicari (misal 'Komang Raim Laode', 'Melompat Lebih Tinggi Sheila on 7').",
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_screenshot",
            "description": (
                "Ambil screenshot tampilan website dan bungkus dalam frame HP iPhone 15 Pro POV (Dynamic Island). "
                "Gunakan saat user minta '/iphone [url]', '/shot [url]', 'screenshot web [url]', atau 'tampilan iphone [url]'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "URL website yang ingin di-screenshot (misal 'https://github.com' atau 'google.com').",
                    },
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "upscale_image",
            "description": (
                "Jernihkan foto buram/pecah dan tingkatkan resolusi foto ke HD (4x Upscale + Sharpening Remini Quality). "
                "Gunakan saat user minta 'hd', 'upscale', 'jernihkan foto', 'perjelas foto', atau 'remini'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "scale": {
                        "type": "integer",
                        "description": "Faktor perbesaran resolusi (default 4x).",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "math_solver",
            "description": (
                "Bantu selesaikan soal matematika, fisika, kimia, atau soal akademis dari foto/teks secara runtut dan jelas. "
                "Gunakan saat user minta '/jawab [soal]', '/soal [soal]', 'bantu jawab soal ini', atau 'selesaikan soal'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {
                        "type": "string",
                        "description": "Teks soal atau deskripsi masalah akademis/matematika yang ingin diselesaikan.",
                    },
                },
                "required": ["question"],
            },
        },
    },
]


# ─── Tool Implementations ─────────────────────────────────────────────────────

USER_LAST_FILES: dict[int, dict] = {}
PENDING_CONVERTED_FILES: dict[int, dict] = {}


def set_user_last_file(user_id: int, file_path: str, filename: str, ext: str) -> None:
    import time
    USER_LAST_FILES[user_id] = {
        "path": file_path,
        "filename": filename,
        "ext": ext.lower().lstrip("."),
        "time": time.time(),
    }


def pop_pending_converted_file(user_id: int) -> dict | None:
    return PENDING_CONVERTED_FILES.pop(user_id, None)


def _tool_file_convert(user_id: int, target_format: str) -> str:
    import os
    import subprocess
    import time

    target_format = target_format.lower().lstrip(".")
    log.info(f"[tool:file_convert] user={user_id} target_format={target_format}")

    file_info = USER_LAST_FILES.get(user_id)
    if not file_info or not os.path.exists(file_info["path"]):
        return json.dumps({
            "error": "Belum ada file yang kamu kirim. Silakan kirim file dulu (DOCX, PDF, JPG, PNG, TXT), baru minta konversi."
        })

    src_path = file_info["path"]
    src_filename = file_info["filename"]
    src_ext = file_info["ext"]

    if src_ext == target_format:
        return json.dumps({
            "error": f"File '{src_filename}' sudah dalam format {target_format.upper()}."
        })

    base_name = os.path.splitext(src_filename)[0]
    out_filename = f"{base_name}.{target_format}"
    temp_dir = os.path.dirname(src_path) or "temp_files"
    out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")

    try:
        # 1. Image -> PDF
        if src_ext in ["jpg", "jpeg", "png", "webp", "bmp"] and target_format == "pdf":
            from PIL import Image
            img = Image.open(src_path)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(out_path, "PDF")

        # 2. Image -> Image (JPG / PNG)
        elif src_ext in ["jpg", "jpeg", "png", "webp", "bmp"] and target_format in ["jpg", "jpeg", "png"]:
            from PIL import Image
            img = Image.open(src_path)
            fmt = "JPEG" if target_format in ["jpg", "jpeg"] else "PNG"
            if fmt == "JPEG" and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(out_path, fmt)

        # 3. PDF -> Image (JPG / PNG)
        elif src_ext == "pdf" and target_format in ["jpg", "jpeg", "png"]:
            import fitz
            doc = fitz.open(src_path)
            if len(doc) == 0:
                return json.dumps({"error": "File PDF kosong."})
            page = doc[0]
            pix = page.get_pixmap(dpi=150)
            pix.save(out_path)
            doc.close()

        # 4. PDF -> DOCX
        elif src_ext == "pdf" and target_format == "docx":
            from pdf2docx import Converter
            cv = Converter(src_path)
            cv.convert(out_path, start=0, end=None)
            cv.close()

        # 5. DOCX -> PDF
        elif src_ext in ["docx", "doc"] and target_format == "pdf":
            import docx
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet

            doc = docx.Document(src_path)
            styles = getSampleStyleSheet()
            doc_pdf = SimpleDocTemplate(out_path, pagesize=letter)
            story = []

            for p in doc.paragraphs:
                txt = p.text.strip()
                if txt:
                    txt_clean = txt.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(txt_clean, styles['Normal']))
                    story.append(Spacer(1, 6))

            if story:
                doc_pdf.build(story)
            else:
                return json.dumps({"error": "Dokumen DOCX tidak berisi teks yang dapat dikonversi ke PDF."})

        # 6. PDF / DOCX / Image -> TXT
        elif target_format == "txt":
            full_text = ""
            if src_ext == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(src_path)
                full_text = "\n\n".join([p.extract_text() or "" for p in reader.pages])
            elif src_ext in ["docx", "doc"]:
                import docx
                doc = docx.Document(src_path)
                full_text = "\n".join([p.text for p in doc.paragraphs])
            else:
                full_text = f"Teks dari file {src_filename}"

            with open(out_path, "w", encoding="utf-8") as f:
                f.write(full_text if full_text.strip() else "Tidak ada teks yang dapat diekstrak.")

        else:
            return json.dumps({
                "error": f"Konversi dari format .{src_ext} ke .{target_format} belum didukung."
            })

        if not os.path.exists(out_path):
            return json.dumps({"error": f"Gagal membuat file hasil konversi {target_format.upper()}."})

        PENDING_CONVERTED_FILES[user_id] = {
            "out_path": out_path,
            "out_filename": out_filename,
            "src_path": src_path,
        }

        return json.dumps({
            "ok": True,
            "out_filename": out_filename,
            "message": f"Konversi '{src_filename}' ke '{out_filename}' berhasil!",
        })

    except Exception as e:
        log.error(f"[file_convert] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal mengonversi file: {e}"})


def _search_ddg_html(query: str, max_results: int = 5) -> list[dict]:
    import urllib.request
    import urllib.parse
    import re
    import ssl

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "id-ID,id;q=0.9,en-US;q=0.8,en;q=0.7",
    }
    url = "https://html.duckduckgo.com/html/?q=" + urllib.parse.quote(query)
    req = urllib.request.Request(url, headers=headers)
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    try:
        with urllib.request.urlopen(req, context=ctx, timeout=2.5) as resp:
            html = resp.read().decode("utf-8", errors="ignore")

        results = []
        snippets = re.findall(r'result__snippet[^>]*>(.*?)</a>', html, re.DOTALL)
        titles = re.findall(r'result__a[^>]*>(.*?)</a>', html, re.DOTALL)

        for t, s in zip(titles, snippets):
            t_clean = re.sub(r'<[^>]+>', '', t).replace('\n', ' ').strip()
            s_clean = re.sub(r'<[^>]+>', '', s).replace('\n', ' ').strip()
            if t_clean and s_clean:
                results.append({"title": t_clean, "snippet": s_clean})
                if len(results) >= max_results:
                    break
        return results
    except Exception as e:
        log.error(f"[html_search] error: {e}")
        return []


def _search_wikipedia(query: str, max_results: int = 3) -> list[dict]:
    import urllib.request
    import urllib.parse
    import json
    import re

    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    url = "https://id.wikipedia.org/w/api.php?action=query&list=search&srsearch=" + urllib.parse.quote(query) + "&format=json"
    req = urllib.request.Request(url, headers=headers)

    try:
        with urllib.request.urlopen(req, timeout=3.0) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            results = []
            for item in data.get("query", {}).get("search", [])[:max_results]:
                title = item.get("title", "")
                snippet = re.sub(r'<[^>]+>', '', item.get("snippet", "")).strip()
                if title and snippet:
                    results.append({"title": title, "snippet": snippet, "url": f"https://id.wikipedia.org/wiki/{urllib.parse.quote(title)}"})
            return results
    except Exception as e:
        log.warning(f"[wiki_search] error: {e}")
        return []


def _tool_web_search(user_id: int, query: str, max_results: int = 5) -> str:
    max_results = min(int(max_results), 10)
    log.info(f"[tool:web_search] user={user_id} query={query!r} n={max_results}")

    results = []

    # 1. Try Wikipedia API (full query or clean core keywords like 'polsub')
    wiki_results = _search_wikipedia(query, max_results=3)
    if not wiki_results and len(query.split()) > 1:
        clean_words = [
            w for w in query.split()
            if w.lower() not in ["daerah", "mana", "apa", "itu", "siapa", "universitas", "kampus", "lokasi", "alamat"]
        ]
        if clean_words:
            wiki_results = _search_wikipedia(" ".join(clean_words), max_results=3)

    if wiki_results:
        results.extend(wiki_results)

    # 2. Try DDGS library
    try:
        with DDGS() as ddgs:
            raw = list(ddgs.text(query, max_results=max_results))
        if raw:
            for r in raw:
                t = r.get("title", "")
                b = r.get("body", "")
                if t and b and not any(kw in t.lower() for kw in ["thesaurus", "antonym", "synonym"]):
                    results.append({"title": t, "snippet": b, "url": r.get("href", "")})
    except Exception as e:
        log.warning(f"[tool:web_search] DDGS library failed ({e})")

    # 3. Try HTML fallback if still empty
    if not results:
        results = _search_ddg_html(query, max_results=max_results)

    if not results:
        return json.dumps({"error": f"Tidak ada hasil ditemukan untuk '{query}'."})

    return json.dumps(results[:max_results], ensure_ascii=False)


import re
from datetime import datetime, timedelta, timezone


def parse_flexible_time(time_str: str) -> datetime | None:
    """Parse relative time strings ('2 detik', '5m', '+30s') or ISO-8601 timestamps."""
    time_str = time_str.strip()

    m = re.search(r'(\d+)\s*(detik|sec|second|s|menit|min|minute|m|jam|hour|h|hari|day|d)', time_str, re.IGNORECASE)
    if m:
        val = int(m.group(1))
        unit = m.group(2).lower()
        now = datetime.now(timezone.utc).replace(tzinfo=None)
        if unit in ['detik', 'sec', 'second', 'seconds', 's']:
            return now + timedelta(seconds=val)
        if unit in ['menit', 'min', 'minute', 'minutes', 'm']:
            return now + timedelta(minutes=val)
        if unit in ['jam', 'hour', 'hours', 'h']:
            return now + timedelta(hours=val)
        if unit in ['hari', 'day', 'days', 'd']:
            return now + timedelta(days=val)

    remind_at_clean = time_str.replace("Z", "").strip()
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"):
        try:
            return datetime.strptime(remind_at_clean, fmt)
        except ValueError:
            pass
    try:
        return datetime.fromisoformat(remind_at_clean)
    except ValueError:
        return None


def _tool_set_reminder(user_id: int, message: str, remind_at: str) -> str:
    log.info(f"[tool:set_reminder] user={user_id} remind_at={remind_at!r}")
    try:
        dt = parse_flexible_time(remind_at)
        if dt is None:
            return json.dumps({"error": f"Format waktu '{remind_at}' tidak dikenali. Gunakan relatif seperti '5 menit', '30 detik', atau ISO format."})
        rid = agent_db.save_reminder(user_id, message, dt)
        return json.dumps({
            "ok": True,
            "reminder_id": rid,
            "remind_at": dt.strftime("%Y-%m-%d %H:%M:%S UTC"),
            "message": message,
        })
    except Exception as e:
        log.error(f"[tool:set_reminder] error: {e}")
        return json.dumps({"error": str(e)})


def _tool_save_note(user_id: int, content: str) -> str:
    log.info(f"[tool:save_note] user={user_id}")
    try:
        note_id = agent_db.save_note(user_id, content)
        return json.dumps({"ok": True, "note_id": note_id, "content": content})
    except Exception as e:
        log.error(f"[tool:save_note] error: {e}")
        return json.dumps({"error": str(e)})


def _tool_get_notes(user_id: int) -> str:
    log.info(f"[tool:get_notes] user={user_id}")
    try:
        notes = agent_db.get_notes(user_id)
        if not notes:
            return json.dumps({"notes": [], "message": "Belum ada catatan tersimpan."})
        return json.dumps({"notes": notes}, ensure_ascii=False, default=str)
    except Exception as e:
        log.error(f"[tool:get_notes] error: {e}")
        return json.dumps({"error": str(e)})


def clean_math_expression(expr: str) -> str:
    """Clean Indonesian math notation like '7 x 7', '7 v 7', '10 : 2', '3,14'."""
    expr = expr.strip()
    # Replace 'x' or 'X' or 'v' or 'V' between numbers or parentheses with '*'
    expr = re.sub(r'(\d+|\))\s*[xXvV]\s*(\d+|\()', r'\1 * \2', expr)
    # Replace ':' between numbers or parentheses with '/'
    expr = re.sub(r'(\d+|\))\s*:\s*(\d+|\()', r'\1 / \2', expr)
    # Replace '^' exponent with '**'
    expr = expr.replace('^', '**')
    # Replace Indonesian comma decimal separator (e.g. 3,14 -> 3.14)
    expr = re.sub(r'(\d+),(\d+)', r'\1.\2', expr)
    return expr


def _tool_calculate(user_id: int, expression: str) -> str:
    log.info(f"[tool:calculate] user={user_id} raw_expr={expression!r}")
    clean_expr = clean_math_expression(expression)
    log.info(f"[tool:calculate] cleaned_expr={clean_expr!r}")

    try:
        import math
        names = {k: getattr(math, k) for k in dir(math) if not k.startswith("_")}
        names["math"] = math
        result = simple_eval(clean_expr, names=names)
        return json.dumps({"expression": expression, "clean_expression": clean_expr, "result": result})
    except InvalidExpression:
        return json.dumps({"error": f"Format hitungan '{expression}' kurang pas. Contoh yang benar: '7 x 7' atau '(10 + 5) * 2'."})
    except ZeroDivisionError:
        return json.dumps({"error": "Pembagian dengan nol tidak diperbolehkan."})
    except Exception as e:
        log.error(f"[tool:calculate] error: {e}")
        return json.dumps({"error": f"Gagal menghitung '{expression}'. Pastikan menggunakan angka dan simbol matematika yang benar."})


def _tool_media_download(user_id: int, url: str, extract_audio: bool = False) -> str:
    """Download video/audio from TikTok, Instagram, YouTube Shorts, Twitter/X using yt_dlp."""
    import yt_dlp
    import time, os

    log.info(f"[tool:media_download] user={user_id} url={url!r} audio={extract_audio}")
    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    out_tmpl = os.path.join(temp_dir, f"media_{user_id}_{int(time.time())}.%(ext)s")

    ydl_opts = {
        "format": "best[filesize<48M]/b[filesize<48M]/bestvideo[filesize<45M]+bestaudio[filesize<45M]/best",
        "outtmpl": out_tmpl,
        "quiet": True,
        "no_warnings": True,
        "max_filesize": 48 * 1024 * 1024,
    }

    try:
        import imageio_ffmpeg
        ydl_opts["ffmpeg_location"] = imageio_ffmpeg.get_ffmpeg_exe()
    except Exception as e:
        log.warning(f"imageio_ffmpeg not available: {e}")

    if extract_audio:
        ydl_opts["format"] = "bestaudio/best"

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            filename = ydl.prepare_filename(info)

            if not os.path.exists(filename):
                matches = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir) if f.startswith(f"media_{user_id}")]
                if matches:
                    filename = matches[0]

            if os.path.exists(filename):
                title = info.get("title", "media")
                out_name = os.path.basename(filename)
                PENDING_CONVERTED_FILES[user_id] = {
                    "out_path": filename,
                    "out_filename": out_name,
                    "src_path": None,
                }
                return json.dumps({
                    "status": "success",
                    "title": title,
                    "filename": out_name,
                    "message": f"Berhasil mengunduh media '{title}'. File sedang dikirim ke chat."
                }, ensure_ascii=False)
            else:
                return json.dumps({"error": "File media gagal diunduh atau ukuran melebihi batas 50 MB."})
    except Exception as e:
        log.error(f"[media_download] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal mengunduh media dari link tersebut: {str(e)}"})


def _tool_remove_bg(user_id: int, bg_color: str = "transparent") -> str:
    """Remove background from user's last uploaded photo and optionally apply background color."""
    import time, os
    log.info(f"[tool:remove_bg] user={user_id} bg_color={bg_color}")
    file_info = USER_LAST_FILES.get(user_id)
    if not file_info or not os.path.exists(file_info["path"]):
        return json.dumps({
            "error": "Belum ada foto yang kamu kirim. Silakan kirim foto dulu, baru minta hapus background."
        })

    src_path = file_info["path"]
    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)

    try:
        from rembg import remove, new_session
        from PIL import Image

        input_img = Image.open(src_path)
        session = new_session("u2netp")
        no_bg = remove(input_img, session=session)

        bg_color_clean = (bg_color or "transparent").lower().strip()
        if bg_color_clean in ["merah", "red"]:
            bg = Image.new("RGBA", no_bg.size, (220, 20, 20, 255))
            bg.paste(no_bg, (0, 0), no_bg)
            out_img = bg.convert("RGB")
            out_filename = "photo_pasfoto_merah.jpg"
            out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
            out_img.save(out_path, "JPEG")
        elif bg_color_clean in ["biru", "blue"]:
            bg = Image.new("RGBA", no_bg.size, (0, 102, 204, 255))
            bg.paste(no_bg, (0, 0), no_bg)
            out_img = bg.convert("RGB")
            out_filename = "photo_pasfoto_biru.jpg"
            out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
            out_img.save(out_path, "JPEG")
        elif bg_color_clean in ["putih", "white"]:
            bg = Image.new("RGBA", no_bg.size, (255, 255, 255, 255))
            bg.paste(no_bg, (0, 0), no_bg)
            out_img = bg.convert("RGB")
            out_filename = "photo_pasfoto_putih.jpg"
            out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
            out_img.save(out_path, "JPEG")
        else:
            out_filename = "photo_nobg.png"
            out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
            no_bg.save(out_path, "PNG")

        PENDING_CONVERTED_FILES[user_id] = {
            "out_path": out_path,
            "out_filename": out_filename,
            "src_path": src_path,
        }
        return json.dumps({
            "status": "success",
            "out_filename": out_filename,
            "message": f"Berhasil menghapus background foto. File '{out_filename}' sedang dikirim..."
        })
    except Exception as e:
        log.error(f"[remove_bg] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal menghapus background foto: {str(e)}"})


def _tool_music_search(user_id: int, query: str) -> str:
    """Search for music/song by title or artist and download high-quality MP3 directly."""
    import yt_dlp
    import time, os

    log.info(f"[tool:music_search] user={user_id} query={query!r}")
    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    out_tmpl = os.path.join(temp_dir, f"music_{user_id}_{int(time.time())}.%(ext)s")

    ydl_opts = {
        "format": "bestaudio/best",
        "outtmpl": out_tmpl,
        "quiet": True,
        "no_warnings": True,
        "default_search": "ytsearch1",
        "max_filesize": 48 * 1024 * 1024,
    }

    try:
        import imageio_ffmpeg
        ydl_opts["ffmpeg_location"] = imageio_ffmpeg.get_ffmpeg_exe()
    except Exception as e:
        log.warning(f"imageio_ffmpeg not available: {e}")

    search_term = f"ytsearch1:{query}"
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(search_term, download=True)
            if "entries" in info and info["entries"]:
                info = info["entries"][0]

            filename = ydl.prepare_filename(info)
            if not os.path.exists(filename):
                matches = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir) if f.startswith(f"music_{user_id}")]
                if matches:
                    filename = matches[0]

            if os.path.exists(filename):
                title = info.get("title", query)
                uploader = info.get("uploader", "Music")
                out_name = f"{title}.mp3"

                PENDING_CONVERTED_FILES[user_id] = {
                    "out_path": filename,
                    "out_filename": out_name,
                    "src_path": None,
                }
                return json.dumps({
                    "status": "success",
                    "title": title,
                    "artist": uploader,
                    "filename": out_name,
                    "message": f"Berhasil menemukan lagu '{title}' oleh {uploader}. File MP3 sedang dikirim ke chat..."
                }, ensure_ascii=False)
            else:
                return json.dumps({"error": f"Lagu '{query}' tidak ditemukan atau gagal diunduh."})
    except Exception as e:
        log.error(f"[music_search] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal mencari/mengunduh lagu: {str(e)}"})


def _create_iphone_mockup(screen_img) -> "Image.Image":
    """Render a clean, native iPhone screen screenshot (Status Bar, Time, 5G, Battery, Dynamic Island, Home Bar)."""
    from PIL import Image, ImageDraw
    import datetime

    target_w, target_h = 390, 844
    screen_resized = screen_img.resize((target_w, target_h), Image.Resampling.LANCZOS).convert("RGB")
    draw = ImageDraw.Draw(screen_resized)

    # 1. Top Status Bar Overlay (Time, Dynamic Island, 5G, Battery)
    now_str = datetime.datetime.now().strftime("%H:%M")

    # Time (Left)
    draw.text((28, 14), now_str, fill="white")

    # Dynamic Island (Center Top)
    island_w, island_h = 115, 30
    island_x = (target_w - island_w) // 2
    island_y = 11
    draw.rounded_rectangle(
        [island_x, island_y, island_x + island_w, island_y + island_h],
        radius=15,
        fill="black"
    )

    # Battery & 5G (Right)
    draw.rounded_rectangle([target_w - 48, 17, target_w - 24, 29], radius=3, outline="white", width=1)
    draw.rectangle([target_w - 46, 19, target_w - 28, 27], fill="white")
    draw.rectangle([target_w - 23, 21, target_w - 22, 25], fill="white")
    draw.text((target_w - 75, 14), "5G", fill="white")

    # 2. Bottom Home Indicator Bar
    bar_w, bar_h = 134, 5
    bar_x = (target_w - bar_w) // 2
    bar_y = target_h - 14
    draw.rounded_rectangle(
        [bar_x, bar_y, bar_x + bar_w, bar_y + bar_h],
        radius=3,
        fill="white"
    )

    # 3. Rounded iOS Screen Corners Mask (radius=44px)
    mask = Image.new("L", (target_w, target_h), 0)
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle([0, 0, target_w, target_h], radius=44, fill=255)

    out_img = Image.new("RGBA", (target_w, target_h), (0, 0, 0, 0))
    out_img.paste(screen_resized, (0, 0), mask)

    return out_img


def _tool_web_screenshot(user_id: int, url: str) -> str:
    """Fetch mobile screenshot of a URL and package it inside an iPhone POV frame."""
    import urllib.request, time, os
    from PIL import Image
    import io

    log.info(f"[tool:web_screenshot] user={user_id} url={url!r}")

    clean_url = url.strip()
    if not clean_url.startswith(("http://", "https://")):
        clean_url = "https://" + clean_url

    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)

    screen_img = None
    microlink_url = f"https://api.microlink.io/?url={clean_url}&screenshot=true&embed=screenshot.url&viewport.width=390&viewport.height=844&viewport.isMobile=true&viewport.deviceScaleFactor=2"

    try:
        req = urllib.request.Request(microlink_url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
        with urllib.request.urlopen(req, timeout=12) as resp:
            img_bytes = resp.read()
            screen_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    except Exception as e:
        log.warning(f"[web_screenshot] Microlink error: {e}, trying thum.io fallback...")
        try:
            thum_url = f"https://image.thum.io/get/width/390/crop/844/mobile/{clean_url}"
            req2 = urllib.request.Request(thum_url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req2, timeout=12) as resp2:
                img_bytes = resp2.read()
                screen_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        except Exception as e2:
            log.error(f"[web_screenshot] thum.io error: {e2}")

    if not screen_img:
        return json.dumps({"error": f"Gagal mengambil screenshot dari website {clean_url}. Pastikan URL dapat diakses."})

    try:
        iphone_mockup = _create_iphone_mockup(screen_img)
        out_filename = "iphone_pov_screenshot.png"
        out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
        iphone_mockup.save(out_path, "PNG")

        PENDING_CONVERTED_FILES[user_id] = {
            "out_path": out_path,
            "out_filename": out_filename,
            "src_path": None,
        }
        return json.dumps({
            "status": "success",
            "out_filename": out_filename,
            "url": clean_url,
            "message": f"Berhasil membuat screenshot iPhone POV dari {clean_url}. Gambar sedang dikirim..."
        }, ensure_ascii=False)
    except Exception as e:
        log.error(f"[web_screenshot] mockup creation error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal memproses mockup iPhone: {str(e)}"})


def _tool_upscale_image(user_id: int, scale: int = 4) -> str:
    """Enhance and upscale user's last uploaded photo to HD Remini-style quality."""
    import time, os
    from PIL import Image, ImageEnhance, ImageFilter

    log.info(f"[tool:upscale_image] user={user_id} scale={scale}")
    file_info = USER_LAST_FILES.get(user_id)
    if not file_info or not os.path.exists(file_info["path"]):
        return json.dumps({
            "error": "Belum ada foto yang kamu kirim. Silakan kirim foto dulu, baru minta jernihkan / upscale HD."
        })

    src_path = file_info["path"]
    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)

    try:
        input_img = Image.open(src_path).convert("RGB")
        w, h = input_img.size

        scale_factor = scale if scale in [2, 3, 4] else 4
        new_w = min(w * scale_factor, 4000)
        new_h = min(h * scale_factor, 4000)

        # 1. High Quality 4x Lanczos Resampling
        hd_img = input_img.resize((new_w, new_h), Image.Resampling.LANCZOS)

        # 2. Unsharp Mask Sharpening
        hd_img = hd_img.filter(ImageFilter.UnsharpMask(radius=2, percent=160, threshold=3))

        # 3. Color & Contrast Enhancement
        color_enhancer = ImageEnhance.Color(hd_img)
        hd_img = color_enhancer.enhance(1.12)

        contrast_enhancer = ImageEnhance.Contrast(hd_img)
        hd_img = contrast_enhancer.enhance(1.08)

        sharpness_enhancer = ImageEnhance.Sharpness(hd_img)
        hd_img = sharpness_enhancer.enhance(1.15)

        out_filename = "photo_hd_remini.png"
        out_path = os.path.join(temp_dir, f"out_{user_id}_{int(time.time())}_{out_filename}")
        hd_img.save(out_path, "PNG", compress_level=6)

        PENDING_CONVERTED_FILES[user_id] = {
            "out_path": out_path,
            "out_filename": out_filename,
            "src_path": src_path,
        }
        return json.dumps({
            "status": "success",
            "out_filename": out_filename,
            "original_size": f"{w}x{h}",
            "hd_size": f"{new_w}x{new_h}",
            "message": f"Berhasil menjernihkan foto ({w}x{h} ➔ {new_w}x{new_h} HD). File HD sedang dikirim..."
        }, ensure_ascii=False)
    except Exception as e:
        log.error(f"[upscale_image] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal menjernihkan foto: {str(e)}"})


def _tool_math_solver(user_id: int, question: str = "") -> str:
    """Solve math, physics, chemistry, or academic problems step-by-step."""
    import os, base64
    log.info(f"[tool:math_solver] user={user_id} question={question!r}")

    file_info = USER_LAST_FILES.get(user_id)
    has_photo = file_info and os.path.exists(file_info["path"])

    solver_prompt = (
        "Kamu adalah Tutor AI Jenius yang sangat ahli Matematika, Fisika, Kimia, dan Soal Akademis. "
        "Tugasmu adalah menganalisis dan menyelesaikan soal dengan format super rapi:\n"
        "1. 📝 **Identifikasi Soal:** Tuliskan ulang soal dengan jelas.\n"
        "2. 💡 **Rumus / Konsep yang Digunakan:** Sebutkan rumus utama yang dipakai.\n"
        "3. 🔍 **Langkah-Langkah Penyelesaian:** Berikan penjelasan runtut, logis, dan gampang dipahami.\n"
        "4. ✅ **Jawaban Akhir:** Highlight hasil/jawaban akhir dengan tebal.\n\n"
        f"Soal/Pertanyaan User: {question or 'Tolong jawab dan selesaikan soal yang ada pada foto ini.'}"
    )

    try:
        from groq import Groq
        api_key = os.environ.get("GROQ_API_KEY")
        if not api_key:
            return json.dumps({"error": "GROQ_API_KEY belum dikonfigurasi di environment."})
        client = Groq(api_key=api_key)

        if has_photo:
            with open(file_info["path"], "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode()

            resp = client.chat.completions.create(
                model="llama-3.2-11b-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": solver_prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}}
                        ]
                    }
                ],
                temperature=0.2,
                max_tokens=1500,
            )
            answer = resp.choices[0].message.content
        else:
            resp = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": "Kamu adalah Tutor AI Jenius matematika & sains."},
                    {"role": "user", "content": solver_prompt}
                ],
                temperature=0.2,
                max_tokens=1500,
            )
            answer = resp.choices[0].message.content

        return json.dumps({
            "status": "success",
            "answer": answer,
        }, ensure_ascii=False)
    except Exception as e:
        log.error(f"[math_solver] error: {e}", exc_info=True)
        return json.dumps({"error": f"Gagal menyelesaikan soal: {str(e)}"})


# ─── Dispatcher ───────────────────────────────────────────────────────────────

_TOOL_MAP = {
    "web_search": _tool_web_search,
    "set_reminder": _tool_set_reminder,
    "save_note": _tool_save_note,
    "get_notes": _tool_get_notes,
    "calculate": _tool_calculate,
    "file_convert": _tool_file_convert,
    "media_download": _tool_media_download,
    "remove_bg": _tool_remove_bg,
    "music_search": _tool_music_search,
    "web_screenshot": _tool_web_screenshot,
    "upscale_image": _tool_upscale_image,
    "math_solver": _tool_math_solver,
}


def execute_tool(user_id: int, tool_name: str, tool_args: dict) -> str:
    """Run a tool by name and return its JSON string result."""
    fn = _TOOL_MAP.get(tool_name)
    if fn is None:
        return json.dumps({"error": f"Tool '{tool_name}' tidak dikenal."})
    return fn(user_id=user_id, **tool_args)
